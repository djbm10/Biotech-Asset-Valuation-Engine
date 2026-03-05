"""
Export utilities: Markdown → DOCX / PDF.

Requires:
  python-docx  (for DOCX export)
  markdown     (for HTML intermediate; `pip install markdown`)

DOCX export uses a simple paragraph-by-paragraph approach that handles
headings, bold, tables, and bullet lists from Markdown.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from bve.valuation.outputs import ValuationOutput


def markdown_to_docx(
    markdown_text: str,
    output_path: str | Path,
    title: str = "Valuation Memo",
    author: str = "BVE Analytics",
) -> Path:
    """
    Convert a Markdown string to a .docx file using python-docx.

    This is a lightweight converter; for production-quality formatting
    consider pandoc (subprocess call) or a paid Markdown→DOCX service.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")

    doc = Document()

    # Styles
    doc.core_properties.author = author
    doc.core_properties.title = title

    lines = markdown_text.split("\n")

    def _add_bold_italic(para, text: str):
        """Parse **bold** and *italic* inline markers."""
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                para.add_run(part)

    in_table = False
    table_rows: list[list[str]] = []

    def _flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        # Skip separator rows (|---|---|)
        data_rows = [r for r in table_rows if not all(re.match(r"[-:]+", c.strip()) for c in r)]
        if len(data_rows) < 2:
            table_rows = []
            in_table = False
            return
        ncols = max(len(r) for r in data_rows)
        tbl = doc.add_table(rows=len(data_rows), cols=ncols)
        tbl.style = "Table Grid"
        for ri, row in enumerate(data_rows):
            for ci, cell_text in enumerate(row[:ncols]):
                cell = tbl.cell(ri, ci)
                cell.text = cell_text.strip()
                if ri == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        doc.add_paragraph("")
        table_rows = []
        in_table = False

    for line in lines:
        stripped = line.strip()

        # Table detection
        if stripped.startswith("|"):
            in_table = True
            cols = [c for c in stripped.split("|") if c.strip() != ""]
            table_rows.append(cols)
            continue
        elif in_table:
            _flush_table()

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        # Horizontal rule
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 60)
        # Bullet
        elif stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _add_bold_italic(para, stripped[2:])
        # Numbered list
        elif re.match(r"^\d+\. ", stripped):
            para = doc.add_paragraph(style="List Number")
            _add_bold_italic(para, re.sub(r"^\d+\. ", "", stripped))
        # Blank line
        elif stripped == "":
            doc.add_paragraph("")
        else:
            # Regular paragraph
            para = doc.add_paragraph()
            _add_bold_italic(para, stripped)

    if in_table:
        _flush_table()

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def export_full_package(
    output: "ValuationOutput",
    memo_type: str = "bd",
    output_dir: str | Path = "memos",
    save_charts: bool = True,
) -> dict[str, str]:
    """
    Generate the complete memo package:
      - Markdown memo
      - DOCX memo
      - All charts (PNG)

    Returns dict of {artifact: path}.
    """
    from bve.reporting.memo_generator import save_memo
    from bve.reporting.charts import save_all_charts

    output_dir = Path(output_dir)
    artifacts: dict[str, str] = {}

    # Markdown
    md_path = save_memo(output, memo_type=memo_type, output_dir=output_dir)
    artifacts["memo_markdown"] = str(md_path)

    # DOCX
    try:
        docx_path = md_path.with_suffix(".docx")
        markdown_to_docx(
            md_path.read_text(encoding="utf-8"),
            docx_path,
            title=f"{output.asset.name} — {memo_type.upper()} Memo",
        )
        artifacts["memo_docx"] = str(docx_path)
    except Exception as e:
        artifacts["memo_docx"] = f"ERROR: {e}"

    # Charts
    if save_charts:
        chart_dir = output_dir / "charts"
        chart_paths = save_all_charts(output, str(chart_dir))
        artifacts.update({f"chart_{k}": v for k, v in chart_paths.items()})

    return artifacts
